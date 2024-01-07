from openai import OpenAI
from dotenv import load_dotenv
import pandas as pd 
from docx import Document

load_dotenv()
client = OpenAI()

def generate_report(file):
  # response = client.chat.completions.create(
  #   model = "gpt-4-1106-preview", 
  #   messages = [
  #     {"role":"system", "content":"You will scan this file, and generate a report for your client. For now, just \
  #       write a summary of the file."},
  #     {"role":"system", "content":f"Here is the file: {pd.read_excel(file)}"},
  #     ],
  #   temperature = 0,
  #   response_format="json",
  # )

  # doc = Document()
  # doc.add_heading(f"Report for : {file.name.split('.')[0]}")
  # doc.add_paragraph(response.choices[0].message.content)
  # doc.save("report.docx")
  # return "report.docx"

 # For testing purposes
  doc = Document()
  doc.add_heading(f"Report: {file.name.split('.')[0]}")
  doc.add_paragraph("This is a test report.")
  doc.save("output/report.docx")
  return 'output/report.docx'